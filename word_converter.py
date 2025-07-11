from docx import Document
from pdf_handler import extract_text_from_pdf

def convert_pdf_to_docx(pdf_path):
    text = extract_text_from_pdf(pdf_path)
    doc = Document()
    doc.add_heading("Converted from PDF", 0)
    doc.add_paragraph(text)
    output_path = "converted.docx"
    doc.save(output_path)
    return output_path
